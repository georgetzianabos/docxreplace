import logging
import re

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

LOGGER = logging.getLogger(__name__)


def search_replace_files(filenames, find, replace, ignore_case=False):

    flags = re.IGNORECASE if ignore_case else 0

    pattern = re.compile(find, flags)

    def correct_paragraph(paragraph):

        changed = False

        for run in paragraph.runs:

            if pattern.search(run.text):
                run.text = pattern.sub(replace, run.text)
                changed = True

        return changed

    for filename in filenames:

        change_count = 0

        try:

            document = Document(filename)

        except PackageNotFoundError:
            LOGGER.error("Failed to process {}".format(filename))
            continue

        for paragraph in document.paragraphs:
            if correct_paragraph(paragraph):
                change_count += 1

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if correct_paragraph(paragraph):
                            change_count += 1

        if change_count:
            LOGGER.debug("Changed {}".format(filename))
            document.save(filename)
